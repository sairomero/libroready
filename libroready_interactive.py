#!/usr/bin/env python3
"""
LibroReady Interactive Formatter - Gives authors full control
Shows issues, lets them review and select what to fix
"""

import sys
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from ebooklib import epub
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER


class InteractiveFormatter:
    """Interactive formatter with full user control"""

    CHAPTER_PATTERNS = [
        r'^(Chapter|CHAPTER|Capítulo|CAPÍTULO)\s+(\d+|[IVXLCDM]+|one|two|three|four|five|six|seven|eight|nine|ten|uno|dos|tres|cuatro|cinco|seis|siete|ocho|nueve|diez)',
        r'^(Ch|CH|Cap|CAP)\.?\s*(\d+)',
        r'^(\d+|[IVXLCDM]+)\.\s*[A-Z]',
        r'^(Part|PART|Parte|PARTE)\s+(\d+|[IVXLCDM]+)',
        r'^(Prólogo|Epílogo|Prologue|Epilogue|Introduction|Introducción)',
    ]

    def __init__(self, docx_path, output_dir=None):
        self.docx_path = Path(docx_path)
        self.output_dir = Path(output_dir) if output_dir else self.docx_path.parent
        self.book_title = self.docx_path.stem
        self.doc = None
        self.detected_chapters = []
        self.issues = []
        self.selected_fixes = []

    def analyze(self):
        """Analyze document and collect all issues"""
        print(f"\n📚 LibroReady Interactive Formatter")
        print(f"📖 Analyzing: {self.docx_path.name}\n")
        print("=" * 70)

        self.doc = Document(self.docx_path)

        # Detect chapters
        print("\n🔍 Analyzing document structure...")
        self._detect_chapters()

        # Check formatting issues
        self._check_formatting_issues()

        print("\n✅ Analysis complete!")

    def _detect_chapters(self):
        """Detect potential chapter/section breaks"""
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()

            if not text or len(text) > 100:
                continue

            is_chapter = False
            detection_method = None

            # Method 1: Pattern matching
            for pattern in self.CHAPTER_PATTERNS:
                if re.match(pattern, text, re.IGNORECASE):
                    is_chapter = True
                    detection_method = 'pattern'
                    break

            # Method 2: Formatting-based (bold + large + short)
            if not is_chapter and i > 30 and len(text) < 50 and para.runs:
                first_run = para.runs[0]
                if (first_run.bold and
                    first_run.font.size and
                    first_run.font.size.pt >= 20):
                    is_chapter = True
                    detection_method = 'formatting'

            if is_chapter:
                self.detected_chapters.append({
                    'index': i,
                    'text': text,
                    'method': detection_method,
                    'selected': True  # Default: include
                })

    def _check_formatting_issues(self):
        """Check for formatting issues"""
        issues_found = {
            'no_indent': 0,
            'inconsistent_spacing': 0,
            'tabs': 0,
            'images': 0
        }

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check for tabs
            for run in para.runs:
                if '\t' in run.text:
                    issues_found['tabs'] += 1
                    break

            # Check indentation
            if para.paragraph_format.first_line_indent is None:
                if not para.style.name.startswith('Heading'):
                    issues_found['no_indent'] += 1

            # Check line spacing
            if para.paragraph_format.line_spacing is None:
                issues_found['inconsistent_spacing'] += 1

        # Add issues to list
        if issues_found['tabs'] > 0:
            self.issues.append({
                'id': 'fix_tabs',
                'name': 'Remove tab characters',
                'description': f"Found {issues_found['tabs']} paragraphs with tab indentation",
                'detail': 'Tabs don\'t convert well to eBook format. Will replace with proper indentation.',
                'severity': 'critical',
                'selected': True
            })

        if issues_found['no_indent'] > 5:
            self.issues.append({
                'id': 'fix_indent',
                'name': 'Add paragraph indentation',
                'description': f"{issues_found['no_indent']} paragraphs missing first-line indent",
                'detail': 'Will add 0.5" first-line indent to body paragraphs.',
                'severity': 'warning',
                'selected': True
            })

        if issues_found['inconsistent_spacing'] > 10:
            self.issues.append({
                'id': 'fix_spacing',
                'name': 'Fix line spacing',
                'description': f"{issues_found['inconsistent_spacing']} paragraphs with inconsistent spacing",
                'detail': 'Will apply 1.15 line spacing throughout.',
                'severity': 'warning',
                'selected': True
            })

        # Always add chapter heading style
        if self.detected_chapters:
            self.issues.append({
                'id': 'apply_headings',
                'name': 'Apply chapter heading styles',
                'description': f"Apply Heading 1 style to {len(self.detected_chapters)} detected chapters",
                'detail': 'Required for automatic table of contents.',
                'severity': 'critical',
                'selected': True
            })

    def show_review(self):
        """Show interactive review interface"""
        print("\n" + "=" * 70)
        print("📋 REVIEW & CUSTOMIZE")
        print("=" * 70)

        # Show detected chapters
        if self.detected_chapters:
            print(f"\n📑 Detected {len(self.detected_chapters)} Chapters/Sections:")
            print("-" * 70)
            for i, chapter in enumerate(self.detected_chapters[:10], 1):
                marker = "✓" if chapter['selected'] else "○"
                print(f"  {marker} {i}. {chapter['text']}")
            if len(self.detected_chapters) > 10:
                print(f"  ... and {len(self.detected_chapters) - 10} more")

            # Ask to review chapters
            print("\n❓ Review chapters? (y/n):", end=" ")
            response = input().strip().lower()
            if response == 'y':
                self._review_chapters()

        # Show formatting issues
        if self.issues:
            print(f"\n🔧 Found {len(self.issues)} Formatting Issues:")
            print("-" * 70)
            for i, issue in enumerate(self.issues, 1):
                marker = "✓" if issue['selected'] else "○"
                severity = "🚫" if issue['severity'] == 'critical' else "⚠️"
                print(f"  {marker} {severity} {issue['name']}")
                print(f"       └─ {issue['description']}")
                print(f"       └─ {issue['detail']}")
                print()

            # Ask which fixes to apply
            print("❓ Choose which issues to fix:")
            print("   [1] Fix all issues (recommended)")
            print("   [2] Select individual fixes")
            print("   [3] Skip all fixes")
            print("\nChoice (1-3):", end=" ")

            choice = input().strip()
            if choice == '2':
                self._select_fixes()
            elif choice == '3':
                for issue in self.issues:
                    issue['selected'] = False
            # else: keep all selected (default)

    def _review_chapters(self):
        """Let user review and edit detected chapters"""
        print("\n📝 Chapter Review")
        print("-" * 70)
        print("Commands:")
        print("  ✓ [number] - Toggle chapter on/off")
        print("  + Add new chapter")
        print("  - [number] - Remove chapter")
        print("  'done' - Finish review")
        print()

        while True:
            # Show current chapters
            for i, chapter in enumerate(self.detected_chapters, 1):
                marker = "✓" if chapter['selected'] else "○"
                print(f"  {i}. {marker} {chapter['text']}")

            print("\nCommand:", end=" ")
            cmd = input().strip().lower()

            if cmd == 'done':
                break
            elif cmd.startswith('+'):
                print("Enter chapter text:", end=" ")
                text = input().strip()
                print("Paragraph number:", end=" ")
                idx = int(input().strip())
                self.detected_chapters.append({
                    'index': idx,
                    'text': text,
                    'method': 'manual',
                    'selected': True
                })
                print(f"✅ Added: {text}")
            elif cmd.isdigit():
                num = int(cmd) - 1
                if 0 <= num < len(self.detected_chapters):
                    self.detected_chapters[num]['selected'] = not self.detected_chapters[num]['selected']
            print()

    def _select_fixes(self):
        """Let user select which fixes to apply"""
        print("\n🔧 Select Fixes to Apply")
        print("-" * 70)

        for i, issue in enumerate(self.issues, 1):
            print(f"\n{i}. {issue['name']}")
            print(f"   {issue['description']}")
            print(f"   Apply this fix? (y/n):", end=" ")

            response = input().strip().lower()
            issue['selected'] = (response == 'y')

    def apply_fixes(self):
        """Apply selected fixes"""
        print("\n⏳ Applying selected fixes...")

        selected_issues = [i for i in self.issues if i['selected']]
        if not selected_issues:
            print("   No fixes selected - document will be formatted as-is")
            return

        # Apply heading styles to selected chapters
        if any(i['id'] == 'apply_headings' for i in selected_issues):
            self._apply_chapter_styles()

        # Apply other fixes
        for issue in selected_issues:
            if issue['id'] == 'fix_tabs':
                self._fix_tabs()
            elif issue['id'] == 'fix_indent':
                self._fix_indentation()
            elif issue['id'] == 'fix_spacing':
                self._fix_spacing()

        print("✅ Fixes applied!")

    def _apply_chapter_styles(self):
        """Apply Heading 1 style to selected chapters"""
        # Ensure Heading 1 style exists
        try:
            heading_style = self.doc.styles['Heading 1']
        except KeyError:
            heading_style = self.doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)

        heading_style.font.name = 'Garamond'
        heading_style.font.size = Pt(18)
        heading_style.font.bold = True
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)

        # Apply to selected chapters
        for chapter in self.detected_chapters:
            if chapter['selected']:
                para = self.doc.paragraphs[chapter['index']]
                para.style = 'Heading 1'

    def _fix_tabs(self):
        """Remove tab characters"""
        for para in self.doc.paragraphs:
            for run in para.runs:
                run.text = run.text.replace('\t', '')

    def _fix_indentation(self):
        """Add first-line indentation"""
        for para in self.doc.paragraphs:
            if para.text.strip() and not para.style.name.startswith('Heading'):
                para.paragraph_format.first_line_indent = Inches(0.5)

    def _fix_spacing(self):
        """Apply consistent line spacing"""
        for para in self.doc.paragraphs:
            para.paragraph_format.line_spacing = 1.15

    def generate_outputs(self):
        """Generate DOCX, EPUB, and PDF"""
        print("\n⏳ Generating output files...")

        # Save formatted DOCX
        formatted_docx = self.output_dir / f"{self.book_title}_formatted.docx"
        self.doc.save(formatted_docx)
        print(f"✅ DOCX: {formatted_docx.name}")

        # Generate EPUB
        epub_path = self._generate_epub()
        print(f"✅ EPUB: {epub_path.name}")

        # Generate PDF
        pdf_path = self._generate_pdf()
        print(f"✅ PDF: {pdf_path.name}")

        print("\n" + "=" * 70)
        print("🎉 SUCCESS! Your book is ready!")
        print("=" * 70)
        print(f"\n📍 Files saved to: {self.output_dir}")
        print(f"\n📤 Next: Upload {epub_path.name} to Amazon KDP")

        return {
            'docx': formatted_docx,
            'epub': epub_path,
            'pdf': pdf_path
        }

    def _generate_epub(self):
        """Generate EPUB file"""
        book = epub.EpubBook()
        book.set_identifier(f'libroready-{self.book_title}')
        book.set_title(self.book_title)
        book.set_language('en')
        book.add_author('Author')

        epub_chapters = []
        chapter_num = 0
        current_content = []
        selected_chapters = [c for c in self.detected_chapters if c['selected']]

        for para in self.doc.paragraphs:
            text = para.text.strip()

            if not text:
                current_content.append('<p>&nbsp;</p>')
                continue

            if para.style.name == 'Heading 1':
                if current_content:
                    c = epub.EpubHtml(
                        title=selected_chapters[chapter_num-1]['text'] if chapter_num > 0 and chapter_num <= len(selected_chapters) else f'Chapter {chapter_num}',
                        file_name=f'chap_{chapter_num:02d}.xhtml',
                        lang='en'
                    )
                    c.content = '<html><body>' + ''.join(current_content) + '</body></html>'
                    book.add_item(c)
                    epub_chapters.append(c)

                chapter_num += 1
                current_content = [f'<h1>{text}</h1>']
            else:
                current_content.append(f'<p>{text}</p>')

        if current_content:
            c = epub.EpubHtml(
                title=selected_chapters[chapter_num-1]['text'] if chapter_num > 0 and chapter_num <= len(selected_chapters) else f'Chapter {chapter_num}',
                file_name=f'chap_{chapter_num:02d}.xhtml',
                lang='en'
            )
            c.content = '<html><body>' + ''.join(current_content) + '</body></html>'
            book.add_item(c)
            epub_chapters.append(c)

        book.toc = epub_chapters
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = ['nav'] + epub_chapters

        epub_path = self.output_dir / f"{self.book_title}.epub"
        epub.write_epub(epub_path, book)
        return epub_path

    def _generate_pdf(self):
        """Generate PDF file"""
        pdf_path = self.output_dir / f"{self.book_title}_print.pdf"

        pdf_doc = SimpleDocTemplate(
            str(pdf_path),
            pagesize=letter,
            rightMargin=72, leftMargin=72,
            topMargin=72, bottomMargin=72
        )

        styles = getSampleStyleSheet()
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Times-Bold'
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            leading=14,
            alignment=TA_JUSTIFY,
            firstLineIndent=36,
            fontName='Times-Roman'
        )

        story = []
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                story.append(Spacer(1, 0.2*inch))
                continue

            if para.style.name == 'Heading 1':
                if story:
                    story.append(PageBreak())
                story.append(Paragraph(text, heading_style))
                story.append(Spacer(1, 0.3*inch))
            else:
                story.append(Paragraph(text, body_style))
                story.append(Spacer(1, 0.1*inch))

        pdf_doc.build(story)
        return pdf_path


def main():
    import argparse

    parser = argparse.ArgumentParser(description='LibroReady Interactive Formatter')
    parser.add_argument('input', help='Input .docx file')
    parser.add_argument('-o', '--output-dir', help='Output directory')
    parser.add_argument('--auto', action='store_true', help='Skip review, apply all fixes automatically')

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"❌ Error: File not found: {args.input}")
        sys.exit(1)

    try:
        formatter = InteractiveFormatter(args.input, args.output_dir)

        # Analyze
        formatter.analyze()

        if not args.auto:
            # Interactive review
            formatter.show_review()

        # Apply fixes
        formatter.apply_fixes()

        # Generate outputs
        formatter.generate_outputs()

    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
