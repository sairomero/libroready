#!/usr/bin/env python3
"""
LibroReady Formatter - Automatic KDP formatting with EPUB and PDF export
Handles everything automatically: chapter detection, formatting, and export
"""

import sys
import os
import re
from pathlib import Path
import zipfile
import shutil
from defusedxml import ElementTree as DefusedET
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from ebooklib import epub
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER


class LibroReadyFormatter:
    """Complete automatic KDP formatter with EPUB and PDF export"""

    # Chapter detection patterns (English + Spanish)
    CHAPTER_PATTERNS = [
        r'^(Chapter|CHAPTER|Capítulo|CAPÍTULO)\s+(\d+|[IVXLCDM]+|one|two|three|four|five|six|seven|eight|nine|ten|uno|dos|tres|cuatro|cinco|seis|siete|ocho|nueve|diez)',
        r'^(Ch|CH|Cap|CAP)\.?\s*(\d+)',
        r'^(\d+|[IVXLCDM]+)\.\s*[A-Z]',  # "1. Title" or "I. Title"
        r'^(Part|PART|Parte|PARTE)\s+(\d+|[IVXLCDM]+)',
        r'^(Prólogo|Epílogo|Prologue|Epilogue|Introduction|Introducción)',
    ]

    NS = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    }

    def __init__(self, docx_path, output_dir=None):
        self.docx_path = Path(docx_path)
        self.output_dir = Path(output_dir) if output_dir else self.docx_path.parent
        self.book_title = self.docx_path.stem
        self.chapters = []

    def process(self):
        """Main processing pipeline - does everything automatically"""
        print(f"\n📚 LibroReady Formatter")
        print(f"📖 Processing: {self.docx_path.name}\n")
        print("=" * 70)

        # Step 1: Load document
        print("\n⏳ Loading document...")
        doc = Document(self.docx_path)

        # Step 2: Detect chapters
        print("⏳ Detecting chapters...")
        self._detect_chapters(doc)

        # Step 3: Apply formatting
        print("⏳ Applying professional formatting...")
        self._apply_formatting(doc)

        # Step 4: Save formatted DOCX
        formatted_docx = self.output_dir / f"{self.book_title}_formatted.docx"
        doc.save(formatted_docx)
        print(f"✅ Formatted DOCX saved: {formatted_docx.name}")

        # Step 5: Generate EPUB
        print("⏳ Generating EPUB for Amazon KDP...")
        epub_path = self._generate_epub(doc)
        print(f"✅ EPUB saved: {epub_path.name}")

        # Step 6: Generate PDF
        print("⏳ Generating print-ready PDF...")
        pdf_path = self._generate_pdf(doc)
        print(f"✅ PDF saved: {pdf_path.name}")

        # Summary
        print("\n" + "=" * 70)
        print("\n🎉 SUCCESS! Your book is ready for Amazon KDP!")
        print(f"\n📦 Generated files:")
        print(f"   📄 {formatted_docx.name}")
        print(f"   📱 {epub_path.name} (for Kindle eBook)")
        print(f"   📄 {pdf_path.name} (for print book)")
        print(f"\n📍 Location: {self.output_dir}")

        if self.chapters:
            print(f"\n📑 Detected {len(self.chapters)} chapters:")
            for i, chapter in enumerate(self.chapters[:5], 1):
                print(f"   {i}. {chapter}")
            if len(self.chapters) > 5:
                print(f"   ... and {len(self.chapters) - 5} more")

        print("\n📤 Next step: Upload the EPUB to Amazon KDP!")
        print("=" * 70 + "\n")

        return {
            'docx': formatted_docx,
            'epub': epub_path,
            'pdf': pdf_path,
            'chapters': len(self.chapters)
        }

    def _detect_chapters(self, doc):
        """Automatically detect chapter titles"""
        detected_chapters = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # Skip empty or very long paragraphs
            if not text or len(text) > 100:
                continue

            # Method 1: Check against chapter patterns
            is_chapter = False
            for pattern in self.CHAPTER_PATTERNS:
                if re.match(pattern, text, re.IGNORECASE):
                    is_chapter = True
                    break

            if is_chapter:
                detected_chapters.append((i, text, 'pattern'))
                continue

            # Method 2: Detect section headers by formatting
            # Skip first 30 paragraphs (usually title pages)
            if i > 30 and len(text) < 50 and para.runs:
                first_run = para.runs[0]

                # Look for: Bold + Large font (20+pt) + Short text
                if (first_run.bold and
                    first_run.font.size and
                    first_run.font.size.pt >= 20):
                    detected_chapters.append((i, text, 'formatting'))

        # Apply detected chapters
        if detected_chapters:
            for idx, text, method in detected_chapters:
                self.chapters.append(text)
                doc.paragraphs[idx]._element.set('chapter-marker', 'true')
            print(f"   Found {len(detected_chapters)} chapters/sections")
        else:
            print("   No chapters detected - will treat as single document")

    def _apply_formatting(self, doc):
        """Apply KDP-compliant formatting"""

        # Ensure Heading 1 style exists
        try:
            heading_style = doc.styles['Heading 1']
        except KeyError:
            heading_style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)

        # Configure Heading 1 style
        heading_style.font.name = 'Garamond'
        heading_style.font.size = Pt(18)
        heading_style.font.bold = True
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)

        # Process all paragraphs
        for para in doc.paragraphs:
            # Apply Heading 1 to detected chapters
            if para._element.get('chapter-marker') == 'true':
                para.style = 'Heading 1'
                para._element.attrib.pop('chapter-marker')  # Clean up marker
                continue

            # Skip if it's already a heading
            if para.style.name.startswith('Heading'):
                continue

            # Skip if no text
            if not para.text.strip():
                continue

            # Apply body text formatting
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.space_after = Pt(0)

            # Apply font formatting to runs
            for run in para.runs:
                if not run.font.name:
                    run.font.name = 'Garamond'
                if not run.font.size:
                    run.font.size = Pt(11)

    def _generate_epub(self, doc):
        """Generate EPUB file for Amazon KDP"""
        book = epub.EpubBook()

        # Set metadata
        book.set_identifier(f'libroready-{self.book_title}')
        book.set_title(self.book_title)
        book.set_language('en')  # TODO: Auto-detect language
        book.add_author('Author')  # TODO: Extract from document

        # Create chapters
        epub_chapters = []
        chapter_num = 0
        current_content = []

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                current_content.append('<p>&nbsp;</p>')
                continue

            # Check if this is a chapter heading
            if para.style.name == 'Heading 1':
                # Save previous chapter if exists
                if current_content:
                    c = epub.EpubHtml(
                        title=self.chapters[chapter_num-1] if chapter_num > 0 and chapter_num <= len(self.chapters) else f'Chapter {chapter_num}',
                        file_name=f'chap_{chapter_num:02d}.xhtml',
                        lang='en'
                    )
                    c.content = '<html><body>' + ''.join(current_content) + '</body></html>'
                    book.add_item(c)
                    epub_chapters.append(c)

                # Start new chapter
                chapter_num += 1
                current_content = [f'<h1>{text}</h1>']
            else:
                # Add paragraph
                current_content.append(f'<p>{text}</p>')

        # Save last chapter
        if current_content:
            c = epub.EpubHtml(
                title=self.chapters[chapter_num-1] if chapter_num > 0 and chapter_num <= len(self.chapters) else f'Chapter {chapter_num}',
                file_name=f'chap_{chapter_num:02d}.xhtml',
                lang='en'
            )
            c.content = '<html><body>' + ''.join(current_content) + '</body></html>'
            book.add_item(c)
            epub_chapters.append(c)

        # Add table of contents
        book.toc = epub_chapters

        # Add navigation files
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # Set spine
        book.spine = ['nav'] + epub_chapters

        # Write EPUB
        epub_path = self.output_dir / f"{self.book_title}.epub"
        epub.write_epub(epub_path, book)

        return epub_path

    def _generate_pdf(self, doc):
        """Generate print-ready PDF"""
        pdf_path = self.output_dir / f"{self.book_title}_print.pdf"

        # Create PDF document
        pdf_doc = SimpleDocTemplate(
            str(pdf_path),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        # Styles
        styles = getSampleStyleSheet()

        # Custom styles
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=RGBColor(0, 0, 0),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Times-Bold'
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            leading=14,
            alignment=TA_JUSTIFY,
            firstLineIndent=36,  # 0.5 inch
            fontName='Times-Roman'
        )

        # Build content
        story = []

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                story.append(Spacer(1, 0.2*inch))
                continue

            if para.style.name == 'Heading 1':
                # Add page break before chapters (except first)
                if story:
                    story.append(PageBreak())
                story.append(Paragraph(text, heading_style))
                story.append(Spacer(1, 0.3*inch))
            else:
                story.append(Paragraph(text, body_style))
                story.append(Spacer(1, 0.1*inch))

        # Build PDF
        pdf_doc.build(story)

        return pdf_path


def main():
    import argparse

    parser = argparse.ArgumentParser(description='LibroReady Formatter - Automatic KDP formatting')
    parser.add_argument('input', help='Input .docx file')
    parser.add_argument('-o', '--output-dir', help='Output directory (default: same as input)')

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"❌ Error: File not found: {args.input}")
        sys.exit(1)

    if not args.input.endswith('.docx'):
        print(f"❌ Error: Input must be a .docx file")
        sys.exit(1)

    try:
        formatter = LibroReadyFormatter(args.input, args.output_dir)
        formatter.process()
    except Exception as e:
        print(f"\n❌ Error processing document: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
