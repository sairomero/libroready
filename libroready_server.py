#!/usr/bin/env python3
"""
LibroReady Flask Server - Backend API for web interface
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import re
import json
import uuid
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from ebooklib import epub
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from libroready_premium import KeywordResearcher, CategoryRecommender, DescriptionOptimizer, SimpleCoverGenerator

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = Path('uploads')
OUTPUT_FOLDER = Path('outputs')
UPLOAD_FOLDER.mkdir(exist_ok=True)
OUTPUT_FOLDER.mkdir(exist_ok=True)

ALLOWED_EXTENSIONS = {'docx'}

# Store session data (in production, use Redis or database)
sessions = {}


class LibroReadyProcessor:
    """Backend processor for LibroReady"""

    CHAPTER_PATTERNS = [
        r'^(Chapter|CHAPTER|Capítulo|CAPÍTULO)\s+(\d+|[IVXLCDM]+)',
        r'^(Ch|CH|Cap|CAP)\.?\s*(\d+)',
        r'^(\d+|[IVXLCDM]+)\.\s*[A-Z]',
        r'^(Part|PART|Parte|PARTE)\s+(\d+|[IVXLCDM]+)',
        r'^(Prólogo|Epílogo|Prologue|Epilogue|Introduction|Introducción|INTRODUCCIÓN)',
    ]

    def __init__(self, docx_path):
        self.docx_path = Path(docx_path)
        self.doc = Document(docx_path)
        self.book_title = self.docx_path.stem

    def analyze(self):
        """Analyze document and return results"""
        chapters = self._detect_chapters()
        issues = self._detect_issues()

        return {
            'chapters': chapters,
            'issues': issues,
            'stats': {
                'total_paragraphs': len(self.doc.paragraphs),
                'total_words': sum(len(p.text.split()) for p in self.doc.paragraphs),
                'has_images': self._count_images()
            }
        }

    def _detect_chapters(self):
        """Detect chapters/sections"""
        chapters = []

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()

            if not text or len(text) > 100:
                continue

            is_chapter = False
            method = None

            # Pattern matching
            for pattern in self.CHAPTER_PATTERNS:
                if re.match(pattern, text, re.IGNORECASE):
                    is_chapter = True
                    method = 'pattern'
                    break

            # Formatting-based detection
            if not is_chapter and i > 30 and len(text) < 50 and para.runs:
                first_run = para.runs[0]
                if (first_run.bold and
                    first_run.font.size and
                    first_run.font.size.pt >= 20):
                    is_chapter = True
                    method = 'formatting'

            if is_chapter:
                chapters.append({
                    'id': f'chapter_{i}',
                    'index': i,
                    'text': text,
                    'method': method
                })

        return chapters

    def _detect_issues(self):
        """Detect formatting issues"""
        issues = []
        counts = {
            'tabs': 0,
            'no_indent': 0,
            'no_spacing': 0
        }

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check tabs
            for run in para.runs:
                if '\t' in run.text:
                    counts['tabs'] += 1
                    break

            # Check indentation
            if para.paragraph_format.first_line_indent is None:
                if not para.style.name.startswith('Heading'):
                    counts['no_indent'] += 1

            # Check spacing
            if para.paragraph_format.line_spacing is None:
                counts['no_spacing'] += 1

        # Create issue objects
        if counts['tabs'] > 0:
            issues.append({
                'id': 'fix_tabs',
                'name': 'Remove tab characters',
                'description': f"Found {counts['tabs']} paragraphs with tab indentation",
                'detail': 'Tabs don\'t convert well to eBook format',
                'severity': 'critical',
                'count': counts['tabs']
            })

        if counts['no_indent'] > 10:
            issues.append({
                'id': 'fix_indent',
                'name': 'Add paragraph indentation',
                'description': f"{counts['no_indent']} paragraphs missing first-line indent",
                'detail': 'Will add 0.5" first-line indent to body paragraphs',
                'severity': 'warning',
                'count': counts['no_indent']
            })

        if counts['no_spacing'] > 10:
            issues.append({
                'id': 'fix_spacing',
                'name': 'Apply consistent line spacing',
                'description': f"{counts['no_spacing']} paragraphs with inconsistent spacing",
                'detail': 'Will apply 1.15 line spacing throughout',
                'severity': 'warning',
                'count': counts['no_spacing']
            })

        return issues

    def _count_images(self):
        """Count images in document"""
        try:
            return len([rel for rel in self.doc.part.rels.values()
                       if "image" in rel.target_ref])
        except:
            return 0

    def process(self, selected_chapters, selected_fixes, output_dir):
        """Apply fixes and generate files"""
        output_dir = Path(output_dir)

        # Apply selected fixes
        if 'fix_tabs' in selected_fixes:
            self._fix_tabs()

        if 'fix_indent' in selected_fixes:
            self._fix_indentation()

        if 'fix_spacing' in selected_fixes:
            self._fix_spacing()

        # Apply chapter styles
        if selected_chapters:
            self._apply_chapter_styles(selected_chapters)

        # Save formatted DOCX
        docx_path = output_dir / f"{self.book_title}_formatted.docx"
        self.doc.save(docx_path)

        # Generate EPUB
        epub_path = self._generate_epub(output_dir, selected_chapters)

        # Generate PDF
        pdf_path = self._generate_pdf(output_dir)

        return {
            'docx': str(docx_path),
            'epub': str(epub_path),
            'pdf': str(pdf_path)
        }

    def _fix_tabs(self):
        """Remove tabs"""
        for para in self.doc.paragraphs:
            for run in para.runs:
                run.text = run.text.replace('\t', '')

    def _fix_indentation(self):
        """Add indentation"""
        for para in self.doc.paragraphs:
            if para.text.strip() and not para.style.name.startswith('Heading'):
                para.paragraph_format.first_line_indent = Inches(0.5)

    def _fix_spacing(self):
        """Apply spacing"""
        for para in self.doc.paragraphs:
            para.paragraph_format.line_spacing = 1.15

    def _apply_chapter_styles(self, selected_chapters):
        """Apply Heading 1 to selected chapters"""
        try:
            heading_style = self.doc.styles['Heading 1']
        except KeyError:
            heading_style = self.doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)

        heading_style.font.name = 'Garamond'
        heading_style.font.size = Pt(18)
        heading_style.font.bold = True
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)

        chapter_indices = {int(c['index']) for c in selected_chapters}
        for i in chapter_indices:
            if i < len(self.doc.paragraphs):
                self.doc.paragraphs[i].style = 'Heading 1'

    def _generate_epub(self, output_dir, selected_chapters):
        """Generate EPUB"""
        book = epub.EpubBook()
        book.set_identifier(f'libroready-{self.book_title}')
        book.set_title(self.book_title)
        book.set_language('en')
        book.add_author('Author')

        # Create a mapping of chapter indices to their actual names
        chapter_names = {c['index']: c['text'] for c in selected_chapters}

        epub_chapters = []
        chapter_num = 0
        current_content = []
        current_chapter_title = None

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()

            if not text:
                current_content.append('<p>&nbsp;</p>')
                continue

            if para.style.name == 'Heading 1':
                # Save previous chapter if exists
                if current_content:
                    c = epub.EpubHtml(
                        title=current_chapter_title or f'Chapter {chapter_num}',
                        file_name=f'chap_{chapter_num:02d}.xhtml',
                        lang='en'
                    )
                    c.content = '<html><body>' + ''.join(current_content) + '</body></html>'
                    book.add_item(c)
                    epub_chapters.append(c)

                # Start new chapter with actual name
                chapter_num += 1
                current_chapter_title = chapter_names.get(i, text)
                current_content = [f'<h1>{text}</h1>']
            else:
                current_content.append(f'<p>{text}</p>')

        # Save last chapter
        if current_content:
            c = epub.EpubHtml(
                title=current_chapter_title or f'Chapter {chapter_num}',
                file_name=f'chap_{chapter_num:02d}.xhtml',
                lang='en'
            )
            c.content = '<html><body>' + ''.join(current_content) + '</body></html>'
            book.add_item(c)
            epub_chapters.append(c)

        book.toc = epub_chapters
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = ['nav'] + epub_chapters

        epub_path = output_dir / f"{self.book_title}.epub"
        epub.write_epub(epub_path, book)
        return epub_path

    def _generate_pdf(self, output_dir):
        """Generate PDF"""
        pdf_path = output_dir / f"{self.book_title}_print.pdf"

        pdf_doc = SimpleDocTemplate(
            str(pdf_path),
            pagesize=letter,
            rightMargin=72, leftMargin=72,
            topMargin=72, bottomMargin=72
        )

        styles = getSampleStyleSheet()
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Times-Bold'
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            leading=14,
            alignment=TA_JUSTIFY,
            firstLineIndent=36,
            fontName='Times-Roman'
        )

        story = []
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                story.append(Spacer(1, 0.2*inch))
                continue

            if para.style.name == 'Heading 1':
                if story:
                    story.append(PageBreak())
                story.append(Paragraph(text, heading_style))
                story.append(Spacer(1, 0.3*inch))
            else:
                try:
                    story.append(Paragraph(text, body_style))
                    story.append(Spacer(1, 0.1*inch))
                except:
                    pass  # Skip problematic paragraphs

        pdf_doc.build(story)
        return pdf_path


# API Endpoints

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Handle file upload and analysis"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not file.filename.endswith('.docx'):
        return jsonify({'error': 'Only .docx files are supported'}), 400

    # Save file
    session_id = str(uuid.uuid4())
    filename = secure_filename(file.filename)
    filepath = UPLOAD_FOLDER / f"{session_id}_{filename}"
    file.save(filepath)

    # Analyze document
    try:
        processor = LibroReadyProcessor(filepath)
        analysis = processor.analyze()

        # Store session data
        sessions[session_id] = {
            'filepath': str(filepath),
            'filename': filename,
            'analysis': analysis
        }

        return jsonify({
            'session_id': session_id,
            'filename': filename,
            'analysis': analysis
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/process', methods=['POST'])
def process_document():
    """Process document with selected options"""
    data = request.json
    session_id = data.get('session_id')
    selected_chapters = data.get('chapters', [])
    selected_fixes = data.get('fixes', [])

    if session_id not in sessions:
        return jsonify({'error': 'Invalid session'}), 400

    session = sessions[session_id]
    filepath = session['filepath']

    try:
        # Create output directory for this session
        output_dir = OUTPUT_FOLDER / session_id
        output_dir.mkdir(exist_ok=True)

        # Process document
        processor = LibroReadyProcessor(filepath)
        results = processor.process(selected_chapters, selected_fixes, output_dir)

        # Store results in session
        sessions[session_id]['results'] = results

        return jsonify({
            'success': True,
            'files': results
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/download/<session_id>/<file_type>')
def download_file(session_id, file_type):
    """Download generated file"""
    if session_id not in sessions:
        return jsonify({'error': 'Invalid session'}), 400

    session = sessions[session_id]
    if 'results' not in session:
        return jsonify({'error': 'No processed files available'}), 400

    results = session['results']
    if file_type not in results:
        return jsonify({'error': 'File type not found'}), 400

    filepath = results[file_type]
    return send_file(filepath, as_attachment=True)


@app.route('/health')
def health():
    """Health check"""
    return jsonify({'status': 'ok'})


@app.route('/')
def index():
    """Serve the landing page"""
    return send_file('index.html')


@app.route('/app.html')
def app_page():
    """Serve the formatting app"""
    return send_file('app.html')


@app.route('/api/premium/keywords', methods=['POST'])
def generate_keywords():
    """Generate keyword recommendations"""
    data = request.json
    session_id = data.get('session_id')
    title = data.get('title', '')
    description = data.get('description', '')

    if session_id not in sessions:
        return jsonify({'error': 'Invalid session'}), 400

    session = sessions[session_id]
    filepath = session['filepath']

    try:
        keyword_tool = KeywordResearcher(filepath, title, description)
        results = keyword_tool.analyze()
        return jsonify(results)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/premium/categories', methods=['POST'])
def recommend_categories():
    """Recommend BISAC categories"""
    data = request.json
    genre = data.get('genre', 'literary')
    themes = data.get('themes', [])
    title = data.get('title', '')

    try:
        category_tool = CategoryRecommender()
        results = category_tool.recommend(genre, themes, title)
        return jsonify({'categories': results})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/premium/description', methods=['POST'])
def optimize_description():
    """Optimize book description"""
    data = request.json
    description = data.get('description', '')
    genre = data.get('genre', 'literary')
    keywords = data.get('keywords', [])

    try:
        desc_tool = DescriptionOptimizer()
        results = desc_tool.optimize(description, genre, keywords)
        return jsonify(results)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/premium/cover', methods=['POST'])
def generate_cover():
    """Generate cover design"""
    data = request.json
    session_id = data.get('session_id')
    title = data.get('title', '')
    author = data.get('author', '')
    genre = data.get('genre', 'literary')
    subtitle = data.get('subtitle', '')

    if session_id not in sessions:
        return jsonify({'error': 'Invalid session'}), 400

    try:
        # Create output directory
        output_dir = OUTPUT_FOLDER / session_id
        output_dir.mkdir(exist_ok=True)

        # Generate cover
        cover_tool = SimpleCoverGenerator()
        cover_img = cover_tool.generate_cover(title, author, genre, subtitle)

        # Save cover
        cover_path = output_dir / f"{session_id}_cover.png"
        cover_tool.save_cover(cover_img, cover_path)

        # Store in session
        if 'premium' not in sessions[session_id]:
            sessions[session_id]['premium'] = {}
        sessions[session_id]['premium']['cover'] = str(cover_path)

        return jsonify({
            'success': True,
            'cover_url': f'/api/premium/cover/{session_id}'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/premium/cover/<session_id>')
def download_cover(session_id):
    """Download generated cover"""
    if session_id not in sessions:
        return jsonify({'error': 'Invalid session'}), 400

    if 'premium' not in sessions[session_id] or 'cover' not in sessions[session_id]['premium']:
        return jsonify({'error': 'No cover generated'}), 400

    cover_path = sessions[session_id]['premium']['cover']
    return send_file(cover_path, mimetype='image/png')


if __name__ == '__main__':
    print("\n🚀 LibroReady Server Starting...")
    print("📍 Server: http://localhost:8080")
    print("📝 Open your browser and go to: http://localhost:8080\n")
    app.run(debug=True, port=8080)
